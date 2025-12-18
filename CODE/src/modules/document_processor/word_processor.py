"""
Word Processor - Extract text, tables from Word documents (.docx).
"""

import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from .base import (
    BaseProcessor,
    ProcessedDocument,
    DocumentMetadata,
    TextChunk,
    Table,
)


class WordProcessor(BaseProcessor):
    """
    Processor cho Word documents (.docx).

    Features:
    - Extract text from paragraphs
    - Extract tables
    - Preserve heading structure
    - Extract metadata
    """

    def __init__(self, config: Dict = None):
        """
        Initialize Word Processor.

        Args:
            config: Configuration dict with options:
                - extract_tables: Extract tables (default: True)
                - include_headers: Mark headers in text (default: True)
        """
        super().__init__(config)
        self.extract_tables_enabled = self.config.get("extract_tables", True)
        self.include_headers = self.config.get("include_headers", True)

    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]

    def process(self, file_path: str) -> ProcessedDocument:
        """Process Word document"""
        start_time = time.time()

        try:
            self.validate(file_path)
        except Exception as e:
            return self._create_error_result(file_path, e, start_time)

        # Check for .doc (old format)
        if file_path.lower().endswith('.doc'):
            return self._create_error_result(
                file_path,
                ValueError("Old .doc format not supported. Please convert to .docx"),
                start_time
            )

        try:
            from docx import Document
        except ImportError:
            return self._create_error_result(
                file_path,
                ImportError("python-docx not installed. Run: pip install python-docx"),
                start_time
            )

        try:
            doc = Document(file_path)
            text_parts = []
            chunks = []
            tables = []
            char_offset = 0
            chunk_index = 0

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Add heading marker if it's a heading
                if self.include_headers and para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    text = f"{'#' * int(level) if level.isdigit() else '#'} {text}"

                text_parts.append(text)

                # Create chunk
                chunks.append(TextChunk(
                    text=text,
                    start_char=char_offset,
                    end_char=char_offset + len(text),
                    chunk_index=chunk_index,
                    metadata={"style": para.style.name}
                ))
                char_offset += len(text) + 1
                chunk_index += 1

            # Extract tables
            if self.extract_tables_enabled:
                for idx, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        table_data.append(row_data)

                    if table_data:
                        tables.append(Table(
                            data=table_data,
                            headers=table_data[0] if table_data else None,
                            table_index=idx
                        ))

                        # Also add table content to text
                        table_text = self._table_to_text(table_data)
                        text_parts.append(f"\n[Table {idx + 1}]\n{table_text}")

            # Metadata
            metadata = self._extract_word_metadata(doc, file_path)

            full_content = "\n".join(text_parts)

            return ProcessedDocument(
                content=full_content,
                chunks=chunks,
                metadata=metadata,
                source_file=file_path,
                file_type="docx",
                tables=tables,
                processed_at=datetime.now(),
                processing_time=time.time() - start_time,
                processor_version=self.VERSION,
                success=True
            )

        except Exception as e:
            return self._create_error_result(file_path, e, start_time)

    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """Convert table to readable text"""
        lines = []
        for row in table_data:
            lines.append(" | ".join(row))
        return "\n".join(lines)

    def _extract_word_metadata(self, doc, file_path: str) -> DocumentMetadata:
        """Extract Word-specific metadata"""
        base_meta = self.extract_metadata(file_path)

        try:
            core_props = doc.core_properties
            base_meta.author = core_props.author
            base_meta.title = core_props.title
            base_meta.extra = {
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "category": core_props.category,
                "comments": core_props.comments,
            }
        except Exception:
            pass

        # Count sections as pages (approximate)
        base_meta.page_count = len(doc.sections)

        return base_meta
